import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import platform


def create_docx_from_markdown(markdown_content: str, title: str = "Document") -> str:
    """
    Creates a DOCX file from markdown content, preserving {{variable}} placeholders.
    
    Args:
        markdown_content: Markdown text with {{variable}} placeholders
        title: Document title
        
    Returns:
        Path to the created temporary DOCX file
    """
    doc = Document()
    
    # Add document title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Split content into lines
    lines = markdown_content.split('\n')
    
    current_paragraph = None
    
    for line in lines:
        line = line.rstrip()
        
        # Empty line - finish current paragraph
        if not line:
            current_paragraph = None
            continue
        
        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            current_paragraph = None
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            current_paragraph = None
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            current_paragraph = None
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            current_paragraph = None
        # Numbered lists
        elif line[0].isdigit() and line[1:3] in ['. ', ') ']:
            doc.add_paragraph(line[3:], style='List Number')
            current_paragraph = None
        else:
            # Regular text - continue paragraph if exists, else create new
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph(line)
            else:
                current_paragraph.add_run('\n' + line)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name


def fill_docx_template(template_path: str, variables: dict, export_type: str) -> str:
    """
    Fills a DOCX template with variables and optionally converts to PDF.
    
    Args:
        template_path: Path to the DOCX template file
        variables: Dictionary of variable key-value pairs
        export_type: Either "docx" or "pdf"
        
    Returns:
        Path to the generated file
    """
    # Load the template
    doc = Document(template_path)
    
    # Replace variables in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Replace variables in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Replace in headers/footers
    for section in doc.sections:
        # Header
        header = section.header
        for paragraph in header.paragraphs:
            for key, value in variables.items():
                placeholder = "{{" + key + "}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        
        # Footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            for key, value in variables.items():
                placeholder = "{{" + key + "}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
    
    # Save filled DOCX
    filled_docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(filled_docx_path)
    
    # If DOCX is requested, return it
    if export_type == "docx":
        return filled_docx_path
    
    # Convert to PDF if requested
    if export_type == "pdf":
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        
        try:
            # Try using LibreOffice for conversion (best quality)
            convert_docx_to_pdf_libreoffice(filled_docx_path, pdf_path)
        except Exception as e:
            print(f"⚠️ LibreOffice conversion failed, trying alternative: {e}")
            try:
                # Fallback: Try using docx2pdf (Windows only)
                from docx2pdf import convert
                convert(filled_docx_path, pdf_path)
            except Exception as e2:
                print(f"⚠️ docx2pdf failed: {e2}")
                # Last resort: basic PDF generation
                convert_docx_to_pdf_basic(filled_docx_path, pdf_path)
        
        # Clean up the intermediate DOCX
        try:
            os.unlink(filled_docx_path)
        except:
            pass
        
        return pdf_path
    
    return filled_docx_path


def convert_docx_to_pdf_libreoffice(docx_path: str, pdf_path: str):
    """Convert DOCX to PDF using LibreOffice (best quality)"""
    output_dir = os.path.dirname(pdf_path)
    
    # LibreOffice command
    if platform.system() == "Windows":
        libreoffice_cmd = "soffice"
    else:
        libreoffice_cmd = "libreoffice"
    
    cmd = [
        libreoffice_cmd,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path
    ]
    
    subprocess.run(cmd, check=True, capture_output=True)
    
    # LibreOffice creates PDF with same name as input
    generated_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    
    # Move to desired location if different
    if generated_pdf != pdf_path:
        os.rename(generated_pdf, pdf_path)


def convert_docx_to_pdf_basic(docx_path: str, pdf_path: str):
    """Basic DOCX to PDF conversion using ReportLab (fallback)"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    
    doc = Document(docx_path)
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            p = Paragraph(paragraph.text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.2 * inch))
    
    pdf_doc.build(story)